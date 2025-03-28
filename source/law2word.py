import openai  # 或百度星河 API
from docx import Document
from openai import OpenAI
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_response(prompt):
    """
    调用百度星河大模型 API 生成合同内容
    """
    api_key = "c7219f76321b1818e9c0df788868adcd81ec2a51"  # https://aistudio.baidu.com/account/accessToken
    client = OpenAI(
        api_key=api_key,
        base_url="https://aistudio.baidu.com/llm/lmapi/v3"  # 星河社区大模型API服务的BaseURL
    )
    
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model="ernie-4.0-turbo-128k",  # 选择模型
        messages=messages,
        max_tokens=2048,
        temperature=0.35,
        top_p=0.3,
        stream=False  # 不使用流式输出
    )
    
    return response.choices[0].message.content

def generate_contract(user_input):
    """
    生成格式规范的法律合同
    """
    prompt = f"""
    请根据以下信息生成一份标准的法律合同：
    {user_input}
    
    要求：
    1. 合同标题居中、加粗，字号16，使用黑色字体。
    2. 章节标题应加粗，字号14，使用黑色字体。
    3. 合同正文使用统一的字体和颜色，不使用特殊符号，如 * 号。
    4. 适当调整段落间距，使合同易于阅读。
    5. 确保合同结构完整，包括合同编号、双方信息、合作条款、违约责任、争议解决等部分。
    """
    
    return get_response(prompt)

def save_contract_to_word(contract_text, filename="contract.docx"):
    """
    以标准格式将合同文本保存为 Word 文档
    """
    doc = Document()
    
    # 添加标题
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    title_run = title.add_run("技术合作合同")
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    doc.add_paragraph("\n合同编号：[填写编号]")
    doc.add_paragraph("签订日期：[填写日期]\n")
    
    # 添加甲乙双方信息
    section_title = doc.add_paragraph("甲方（技术提供方）：")
    section_title.bold = True
    section_title_run = section_title.runs[0]
    section_title_run.font.size = Pt(14)
    
    doc.add_paragraph("名称：张三\n地址：[张三的具体地址]\n联系方式：[张三的联系电话/邮箱]\n")
    
    section_title = doc.add_paragraph("乙方（市场推广方）：")
    section_title.bold = True
    section_title_run = section_title.runs[0]
    section_title_run.font.size = Pt(14)
    
    doc.add_paragraph("名称：李四\n地址：[李四的具体地址]\n联系方式：[李四的联系电话/邮箱]\n")
    
    # 添加合同正文
    for section in contract_text.split("\n\n"):
        if section.strip():
            paragraph = doc.add_paragraph(section.strip())
            paragraph_run = paragraph.runs[0]
            paragraph_run.font.size = Pt(12)
    
    doc.save(filename)
    print(f"合同已保存为 {filename}")

if __name__ == "__main__":
    user_input = "甲方：张三，乙方：李四，合同内容涉及技术合作，期限1年，甲方负责提供技术支持，乙方负责市场推广..."
    contract_text = generate_contract(user_input)
    save_contract_to_word(contract_text)
